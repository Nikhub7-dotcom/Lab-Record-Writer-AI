from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def save_as_docx(lab_record_text, experiment_name, output_path):
    doc = Document()

    # --- Page title ---
    title = doc.add_heading('', 0)
    title_run = title.add_run(experiment_name.upper())
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # empty line after title

    # --- Parse and write each section ---
    sections = lab_record_text.split('\n')

    for line in sections:
        line = line.strip()
        if not line:
            doc.add_paragraph()  # blank line
            continue

        # section headers like "1. AIM" or "AIM"
        is_header = (
            line.startswith("1.") or line.startswith("2.") or
            line.startswith("3.") or line.startswith("4.") or
            line.startswith("5.") or line.startswith("6.") or
            line.upper() in ["AIM", "THEORY", "PROCEDURE",
                             "CODE", "EXPECTED OUTPUT", "CONCLUSION"]
        )

        if is_header:
            heading = doc.add_paragraph()
            heading_run = heading.add_run(line)
            heading_run.font.bold = True
            heading_run.font.size = Pt(13)
            heading_run.font.color.rgb = RGBColor(0, 0, 128)  # dark blue
        else:
            para = doc.add_paragraph(line)
            para.runs[0].font.size = Pt(11)

    # --- Save the file ---
    doc.save(output_path)
    print(f"Saved: {output_path}")